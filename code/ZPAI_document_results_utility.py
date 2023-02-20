# from asyncio.windows_events import NULL
from logging import NullHandler
from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Pt

from docx.shared import RGBColor

from collections import OrderedDict


import yaml
from pathlib import Path
#import datetime
from datetime import date

import pandas as pd

import matplotlib.pyplot as plt

import seaborn as sns


# load CSV data file
def load_csv_data(csv_path):
    return pd.read_csv(csv_path, delimiter=',', index_col=[0])

# Get the column, index and value for the min value within a dataframe. 
def get_min_values(df):
    min_col_values = df.min()
    min_col_values = min_col_values.to_frame()
    min_column = min_col_values[[0]].idxmin()
    min_index = df[[min_column[0]]].idxmin()
    min_value = df.min().min()

    return(min_column[0], min_index[0], min_value)

# Get the column, index and value for the max value within a dataframe. 
def get_max_values(df):
    max_col_values = df.max()
    max_col_values = max_col_values.to_frame()
    max_column = max_col_values[[0]].idxmax()
    max_index = df[[max_column[0]]].idxmax()
    max_value = df.max().max()

    return(max_column[0], max_index[0], max_value)

def initialize_document(summery_result_values, m_date):
    document = Document()

    # set page size
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    heading_0 = "Measurement overview!"
    document.add_heading(heading_0, 0)
    # document.add_heading("Results for measurement on {}!".format(data_x['measurement_date'], 0))

    document.add_paragraph('Measurement date: {}'.format(m_date), style='List Bullet')
    document.add_paragraph('Random seed: {}'.format(summery_result_values['random_seed']), style='List Bullet')
    document.add_paragraph('Number of PCA calculated features: {}'.format(summery_result_values['pca_numbers']), style='List Bullet')
    document.add_paragraph('Outliert detection per bin: {}'.format(summery_result_values['bin_outlier_detect']), style='List Bullet')

    if 'autosklearn_runtime' in summery_result_values:
        document.add_paragraph('Autosklearn runtime per dataset: {} seconds'.format(summery_result_values['autosklearn_runtime']), style='List Bullet')
        document.add_paragraph('Autosklearn time limit per algorithm: {} seconds'.format(summery_result_values['autosklearn_limit']), style='List Bullet')

    return document

def initialize_section(document, summery_result_values, const_machine_model):
    heading_1 = "Results for {}".format(const_machine_model)
    document.add_heading(heading_1, level=1)

    document.add_paragraph('Input data file name: {}'.format(summery_result_values[const_machine_model]['input_file_name']), style='List Bullet')
    document.add_paragraph('Input data file creation date: {}'.format(summery_result_values[const_machine_model]['input_file_creation_date']), style='List Bullet')
    document.add_paragraph('Number of processed data points: {}'.format(summery_result_values[const_machine_model]['input_file_size']), style='List Bullet')

    return document

def add_results_heading(document, heading, heading_level):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)

    heading_2 = heading
    document.add_heading(heading_2, level=heading_level)

def get_results_values(score, classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results):
    score_result_df = pd.DataFrame()
    
    # Extract value for neural nets.
    if not nn_results.empty:
        nn_score_results = nn_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([nn_score_results])])
        score_result_df = score_result_df.rename(index={0: 'nn'})

    # Extract value for classical algorithms.
    if not classic_results.empty:
        classic_score_results = classic_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([classic_score_results])])
        score_result_df = score_result_df.rename(index={0: 'classic'})


    # Extract value for auto-sklearn method.
    if not autosklearn_results.empty:
        autosklearn_score_results = autosklearn_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([autosklearn_score_results])])
        score_result_df = score_result_df.rename(index={0: 'autosklearn'})

    
    # Extract value for auto-gluon method.
    if not autogluon_results.empty:
        autogluon_score_results = autogluon_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([autogluon_score_results])])
        score_result_df = score_result_df.rename(index={0: 'autogluon'})

    # Extract value for flaml method.
    if not flaml_results.empty:
        flaml_score_results = flaml_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([flaml_score_results])])
        score_result_df = score_result_df.rename(index={0: 'flaml'})

    # Extract value for optuna method.
    if not optuna_results.empty:
        optuna_score_results = optuna_results.loc[ f'Test-{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([optuna_score_results])])
        score_result_df = score_result_df.rename(index={0: 'optuna'})

    score_result_df = score_result_df.astype(float)

    if score=="R2" or score=="MAPE" or score=="N-RMSE" or score=="IQR-RMSE" or score=="CV-RMSE":
        score_result_df = score_result_df.round(decimals=4)
    elif score=="MAE" or score=="RMSE":
        score_result_df = score_result_df.astype(int)

    return score_result_df

def add_score_table(document, score_result_df, score, table_font_size):
    score_col_name_max, score_row_name_max, max_value =  get_max_values(score_result_df)
    score_col_name_min, score_row_name_min, min_value =  get_min_values(score_result_df)

    score_extended_column_names = list(score_result_df.columns.values)
    score_row_names = list(score_result_df.index.values)

    score_extended_column_names.insert(0, ' ')

    score_table = document.add_table(rows=1, cols=len(score_extended_column_names))

    # set table style
    score_table.style = document.styles['Light Shading Accent 1']

    hdr_cells = score_table.rows[0].cells

    for i, val in enumerate(score_extended_column_names):
        hdr_cells[i].text = str(val)

    # add a data row for each item
    for i in range(len(score_row_names)):
        cells = score_table.add_row().cells
        # for j in range(len(r2_extended_column_names)):
        for j, val in enumerate(score_extended_column_names):
            if j == 0:
                cells[j].text = score_row_names[i]
            else:
                cells[j].text = str(score_result_df.iat[i, j-1])

                # set the style - text size & bold 
                paragraphs = cells[j].paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(table_font_size) # set the text size

                if score=="R2":
                    # check for the min & max values and set them to bold
                    if (cells[0].text == str(score_row_name_max)): # check if it is the line with the highest R2-score
                        if(str(val) == str(score_col_name_max)): # check if it is the R2-Score column 
                            font.bold = True # set value to bold
                            font.color.rgb = RGBColor(209, 25, 190) # set the color
                elif score=="MAE" or score=="RMSE" or score=="MAPE" or score=="N-RMSE" or score=="IQR-RMSE" or score=="CV-RMSE" or score=="MEV":
                    # check for the min & max values and set them to bold
                    if (cells[0].text == str(score_row_name_min)): # check if it is the line with the highest R2-score
                        if(str(val) == str(score_col_name_min)): # check if it is the R2-Score column 
                            font.bold = True # set value to bold
                            font.color.rgb = RGBColor(209, 25, 190) # set the color

    for row in score_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(table_font_size)

def add_bar_plot(document, score_result_df, score, measurement_date, input_file_creation_date, const_machine_model, repo_path, picture_size):
    # add R2 chart
    sub_dir = "{}-{}-{}-{}".format(measurement_date, const_machine_model,'final',input_file_creation_date)
    filename = "{}-{}-{}-{}-{}.{}".format(measurement_date, const_machine_model,'final',input_file_creation_date,f'{score}-score', 'png')
    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)

    x, y, score_max_value = get_max_values(score_result_df)
    plt.rcParams["figure.figsize"] = (15,10)
    plt.rcParams.update({'font.size': 15})
    sns.set_style("whitegrid")
    ax = score_result_df.plot(kind="bar", ylim=(0, score_max_value * 1.5), rot=0)
    graphics_title = f"{score}-score for {const_machine_model}"
    # plt.title(graphics_title, fontsize=22) # uncomment for publication
    plt.xlabel("Methods", fontsize=20)
    plt.ylabel(f"{score}-score", fontsize=20)

    ################
    # modify x axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_xticklabels()]
    # Capitalize the first letters of the label names
    labels = [label.capitalize() for label in labels]
    # Replace 'classic' string by the actual algorithm 'Rand. Forest'
    labels = [label.replace('Classic', 'Rand. Forest') for label in labels]
    # replace 'nn' string by 'Neural Networkt' string
    labels = [label.replace('Nn', 'Neural Network') for label in labels]
    # replace 'Autogluon' string by 'AutoGluon' string
    labels = [label.replace('Autogluon', 'AutoGluon') for label in labels]
    # replace 'Autosklearn' string by 'auto-sklearn' string
    labels = [label.replace('Autosklearn', 'auto-sklearn') for label in labels]
    # replace 'Flaml' string by 'FLAML' string
    labels = [label.replace('Flaml', 'FLAML') for label in labels]

    ax.set_xticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = ax.get_legend_handles_labels()
    # Capitalize the first letters
    # labels_txt = [name.capitalize() for name in labels_txt]
    # replace 'classic' string by the actual algorithm 'Rand. Forest'
    labels_txt = [w.replace('extension', 'series') for w in labels_txt]
    # replace 'nn' string by the actual algorithm 'Neural Networkt'
    # labels_txt = [w.replace('Nn', 'Neural Network') for w in labels_txt]

    by_label = OrderedDict(zip(labels_txt, handles))
    # ax.legend(by_label.values(), by_label.keys(), loc='upper center', shadow=True)
    ax.legend(by_label.values(), by_label.keys())


    plt.savefig(Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename))

    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()

def add_line_plot(document, score_result_df, score, measurement_date, input_file_creation_date, const_machine_model, repo_path, picture_size):
    # add R2 chart
    sub_dir = "{}-{}-{}-{}".format(measurement_date, const_machine_model,'final',input_file_creation_date)
    filename = "{}-{}-{}-{}-{}.{}".format(measurement_date, const_machine_model,'final',input_file_creation_date,f'{score}-line-score', 'png')
    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)

    # get the max value for a dynamic adaptation of the y-axes
    # x, y, score_max_value = get_max_values(score_result_df)
    plt.rcParams["figure.figsize"] = (15,10)
    plt.rcParams.update({'font.size': 15})
    sns.set_style("whitegrid")
    score_result_df = score_result_df.transpose()
    ax = score_result_df.plot.line()

    # delete the borderes of the plot
    sns.despine(left=True, bottom=True)
    # set title
    graphics_title = f"{score}-score for {const_machine_model}"
    # plt.title(graphics_title, fontsize=22) # uncomment for publication

    ################
    # modify x axis labels
    #################
    # Extract the x axis label names 
    labels = [item.get_text() for item in ax.get_xticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Replace 'extension' string by 'series'
    labels = [label.replace('extension', 'series') for label in labels]

    ax.set_xticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = ax.get_legend_handles_labels()
    # Capitalize the first letters
    labels_txt = [name.capitalize() for name in labels_txt]
    # replace 'classic' string by the actual algorithm 'Rand. Forest'
    labels_txt = [w.replace('Classic', 'Rand. Forest') for w in labels_txt]
    # replace 'nn' string by the actual algorithm 'Neural Networkt'
    labels_txt = [w.replace('Nn', 'Neural Network') for w in labels_txt]
    # replace 'Autogluon' string by 'AutoGluon' string
    labels_txt = [w.replace('Autogluon', 'AutoGluon') for w in labels_txt]
    # replace 'Autosklearn' string by 'auto-sklearn' string
    labels_txt = [w.replace('Autosklearn', 'auto-sklearn') for w in labels_txt]
    # replace 'Flaml' string by 'FLAML' string
    labels_txt = [w.replace('Flaml', 'FLAML') for w in labels_txt]

    by_label = OrderedDict(zip(labels_txt, handles))
    # ax.legend(by_label.values(), by_label.keys(), loc='upper center', shadow=True)
    ax.legend(by_label.values(), by_label.keys())

    plt.xlabel("Feature combination", fontsize=20)
    plt.ylabel(f"{score}-score", fontsize=20)

    plt.savefig(Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename))

    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()

def add_duration_line_plot(document, score_result_df, name, measurement_date, input_file_creation_date, const_machine_model, repo_path, picture_size):
    # add R2 chart
    sub_dir = "{}-{}-{}-{}".format(measurement_date, const_machine_model,'final',input_file_creation_date)
    filename = "{}-{}-{}-{}-{}.{}".format(measurement_date, const_machine_model,'final',input_file_creation_date,f'{name}', 'png')
    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)

    plt.rcParams["figure.figsize"] = (15,10)
    plt.rcParams.update({'font.size': 15})
    sns.set_style("whitegrid")
    score_result_df = score_result_df.transpose()
    ax = score_result_df.plot.line()
    graphics_title = f"{name} in seconds for {const_machine_model}"
    # plt.title(graphics_title, fontsize=22)

    ################
    # modify x axis labels
    #################
    # Extract the x axis label names 
    labels = [item.get_text() for item in ax.get_xticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Replace 'extension' string by 'series'
    labels = [label.replace('extension', 'series') for label in labels]

    ax.set_xticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = ax.get_legend_handles_labels()
    # Capitalize the first letters
    labels_txt = [name.capitalize() for name in labels_txt]
    # replace 'classic' string by the actual algorithm 'Rand. Forest'
    labels_txt = [w.replace('Classic', 'Rand. Forest') for w in labels_txt]
    # replace 'nn' string by the actual algorithm 'Neural Networkt'
    labels_txt = [w.replace('Nn', 'Neural Network') for w in labels_txt]
    # replace 'Autogluon' string by 'AutoGluon' string
    labels_txt = [w.replace('Autogluon', 'AutoGluon') for w in labels_txt]
    # replace 'Autosklearn' string by 'auto-sklearn' string
    labels_txt = [w.replace('Autosklearn', 'auto-sklearn') for w in labels_txt]
    # replace 'Flaml' string by 'FLAML' string
    labels_txt = [w.replace('Flaml', 'FLAML') for w in labels_txt]

    by_label = OrderedDict(zip(labels_txt, handles))
    # ax.legend(by_label.values(), by_label.keys(), loc='upper center', shadow=True)
    ax.legend(by_label.values(), by_label.keys())

    plt.xlabel("Feature combination", fontsize=18)
    plt.ylabel(f"{name} in seconds", fontsize=18)

    plt.savefig(Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename))

    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()

def add_stats_table(document, best_result_df, summery_result_values, table_font_size):
    extended_column_names = [' ','# data points', 'Mean', 'Median', 'Std. diviation', 'Max', 'Min']
    row_names = list(best_result_df.index.values)
    

    stats_table = document.add_table(rows=1, cols=len(extended_column_names))

    # set table style
    stats_table.style = document.styles['Light Shading Accent 5']

    hdr_cells = stats_table.rows[0].cells

    # set the column names
    for i, val in enumerate(extended_column_names):
        hdr_cells[i].text = str(val)

    # add a data row for each item
    for i in range(len(row_names)):
        cells = stats_table.add_row().cells
        # for j in range(len(extended_column_names)):
        for j, val in enumerate(extended_column_names):
            if j == 0: # set construction machine name
                cells[j].text = row_names[i]
            elif j == 1: # set input file size
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_size'])
            elif j == 2: # set mean
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_mean'])
            elif j == 3: # set median
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_50'])
            elif j == 4: # set stand. deviation
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_std'])
            elif j == 5: # set max value
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_max'])
            elif j == 6: # set min value
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_min'])

    for row in stats_table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(table_font_size)

def add_summary_table(document, best_result_df, summery_result_values, table_font_size):
    best_max_scores = best_result_df.astype('float64').idxmax(axis=0)
    best_min_scores = best_result_df.astype('float64').idxmin(axis=0)

    extended_column_names = list(best_result_df.columns.values)
    row_names = list(best_result_df.index.values)

    extended_column_names.insert(0, ' ')
    extended_column_names.insert(1, '# data points')
    
    table = document.add_table(rows=1, cols=len(extended_column_names))

    # set table style
    table.style = document.styles['Light Shading Accent 5']

    hdr_cells = table.rows[0].cells

    # set the column names
    for i, val in enumerate(extended_column_names):
        hdr_cells[i].text = str(val)

    # add a data row for each item
    for i in range(len(row_names)):
        cells = table.add_row().cells
        # for j in range(len(extended_column_names)):
        for j, val in enumerate(extended_column_names):
            if j == 0: # set construction machine name
                cells[j].text = row_names[i]
            elif j == 1: # set input file size
                cells[j].text = str(summery_result_values[str(row_names[i])]['input_file_size'])
            else: # set values for R2, MAE & RMSE
                cells[j].text = str(best_result_df.iat[i, j-2])

                # set the style - text size & bold 
                paragraphs = cells[j].paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(table_font_size) # set the text size

                # check for the min & max values and set them to bold
                if (cells[0].text == str(best_max_scores['R2-Score'])): # check if it is the line with the highest R2-score
                    if(str(val) == 'R2-Score'): # check if it is the R2-Score column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color
                
                if (cells[0].text == str(best_min_scores['MAPE'])): # check if it is the line with the lowest score
                    if(str(val) == 'MAPE'): # check if it is the MAPE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

                if (cells[0].text == str(best_min_scores['MAE'])): # check if it is the line with the lowest score
                    if(str(val) == 'MAE'): # check if it is the MAE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

                if (cells[0].text == str(best_min_scores['RMSE'])): # check if it is the line with the lowest score
                    if(str(val) == 'RMSE'): # check if it is the RMSE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

                if (cells[0].text == str(best_min_scores['N-RMSE'])): # check if it is the line with the lowest score
                    if(str(val) == 'N-RMSE'): # check if it is the N-RMSE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

                if (cells[0].text == str(best_min_scores['IQR-RMSE'])): # check if it is the line with the lowest score
                    if(str(val) == 'IQR-RMSE'): # check if it is the N-RMSE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

                if (cells[0].text == str(best_min_scores['CV-RMSE'])): # check if it is the line with the lowest score
                    if(str(val) == 'CV-RMSE'): # check if it is the N-RMSE column 
                        font.bold = True # set value to bold
                        font.color.rgb = RGBColor(209, 25, 190) # set the color

    for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(table_font_size)

def get_duration_results_values(score, classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results):
    score_result_df = pd.DataFrame()
    
    # Extract value for neural nets.
    if not nn_results.empty:
        nn_score_results = nn_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([nn_score_results])])
        score_result_df = score_result_df.rename(index={0: 'nn'})

    # Extract value for classical algorithms.
    if not classic_results.empty:
        classic_score_results = classic_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([classic_score_results])])
        score_result_df = score_result_df.rename(index={0: 'classic'})


    # Extract value for auto-sklearn method.
    if not autosklearn_results.empty:
        autosklearn_score_results = autosklearn_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([autosklearn_score_results])])
        score_result_df = score_result_df.rename(index={0: 'autosklearn'})

    
    # Extract value for auto-gluon method.
    if not autogluon_results.empty:
        autogluon_score_results = autogluon_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([autogluon_score_results])])
        score_result_df = score_result_df.rename(index={0: 'autogluon'})

    # Extract value for flaml method.
    if not flaml_results.empty:
        flaml_score_results = flaml_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([flaml_score_results])])
        score_result_df = score_result_df.rename(index={0: 'flaml'})

    # Extract value for optuna method.
    if not optuna_results.empty:
        optuna_score_results = optuna_results.loc[ f'{score}' , : ]
        score_result_df = pd.concat([score_result_df, pd.DataFrame.from_records([optuna_score_results])])
        score_result_df = score_result_df.rename(index={0: 'optuna'})

    score_result_df = score_result_df.astype(float)

    if score=="Duration":
        score_result_df = score_result_df.round(decimals=4)

    return score_result_df


def add_box_plot(document, score_result_df, score, measurement_date, input_file_creation_date, const_machine_model, repo_path, picture_size):

    sub_dir = "{}-{}-{}-{}".format(measurement_date, const_machine_model,'final',input_file_creation_date)
    filename = "{}-{}-{}-{}-{}.{}".format(measurement_date, const_machine_model,'final',input_file_creation_date,f'{score}-boxplot', 'png')
    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)

    # get the max value for a dynamic adaptation of the y-axes
    # x, y, score_max_value = get_max_values(score_result_df)
    plt.rcParams["figure.figsize"] = (15,10)
    plt.rcParams.update({'font.size': 15})
    sns.set_style("whitegrid")
    # score_result_df = score_result_df.transpose()
    # print(score_result_df)
    sns.boxenplot(data = score_result_df['extension-location'])
    # ax = score_result_df.plot.line()

    # delete the borderes of the plot
    # sns.despine(left=True, bottom=True)
    # set title
    # graphics_title = f"{score}-score for {const_machine_model}"
    # plt.title(graphics_title, fontsize=22) # uncomment for publication

    ################
    # modify x axis labels
    #################
    # Extract the x axis label names 
    # labels = [item.get_text() for item in ax.get_xticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Replace 'extension' string by 'series'
    # labels = [label.replace('extension', 'series') for label in labels]

    # ax.set_xticklabels(labels)

    ################
    # modify legend
    #################
    # handles, labels_txt = ax.get_legend_handles_labels()
    # # Capitalize the first letters
    # labels_txt = [name.capitalize() for name in labels_txt]
    # # replace 'classic' string by the actual algorithm 'Rand. Forest'
    # labels_txt = [w.replace('Classic', 'Rand. Forest') for w in labels_txt]
    # # replace 'nn' string by the actual algorithm 'Neural Networkt'
    # labels_txt = [w.replace('Nn', 'Neural Network') for w in labels_txt]

    # by_label = OrderedDict(zip(labels_txt, handles))
    # # ax.legend(by_label.values(), by_label.keys(), loc='upper center', shadow=True)
    # ax.legend(by_label.values(), by_label.keys())

    # plt.xlabel("Feature combination", fontsize=20)
    # plt.ylabel(f"{score}-score", fontsize=20)

    # save grafics
    plt.savefig(CHART_PATH)
    # plt.show()
    # integrate grafics into dokumentation
    # document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    # plt.savefig(Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename))

    # CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
    # document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()

def add_box_plot2(document, score_result_df, column_name, score, measurement_date, input_file_creation_date, const_machine_model, repo_path, picture_size):
    # add R2 chart
    sub_dir = "{}-{}-{}-{}".format(measurement_date, const_machine_model,'final',input_file_creation_date)
    filename = "{}-{}-{}-{}-{}.{}".format(measurement_date, const_machine_model,'final',input_file_creation_date,f'{score}-boxplot', 'png')
    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)

    # x, y, score_max_value = get_max_values(score_result_df)
    plt.rcParams["figure.figsize"] = (15,10)
    plt.rcParams.update({'font.size': 15})
    sns.set_style("whitegrid")
    # print(score_result_df)
    ax = sns.boxenplot(data = score_result_df[column_name])
    graphics_title = f"{score}-score for {const_machine_model}"
    # plt.title(graphics_title, fontsize=22) # uncomment for publication
    plt.xlabel("Methods", fontsize=20)
    plt.ylabel("Mean absolute percentage error", fontsize=20)

    ################
    # modify x axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_xticklabels()]
    # print(labels)
    # Capitalize the first letters of the label names
    labels = [label.capitalize() for label in labels]
    # Replace 'classic' string by the actual algorithm 'Rand. Forest'
    labels = [label.replace('0', 'Autogluon') for label in labels]
    labels = [label.replace('1', 'auto-sklearn') for label in labels]
    labels = [label.replace('2', 'Rand. Forest') for label in labels]
    labels = [label.replace('3', 'FLAML') for label in labels]
    labels = [label.replace('4', 'Neural Net') for label in labels]
    labels = [label.replace('5', 'Optuna') for label in labels]

    ax.set_xticklabels(labels)


    plt.savefig(Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename))

    CHART_PATH = Path(repo_path, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()


def add_simple_score_table(document, score_result_df, table_font_size):

    score_extended_column_names = list(score_result_df.columns.values)
    score_row_names = list(score_result_df.index.values)

    score_extended_column_names.insert(0, ' ')

    score_table = document.add_table(rows=1, cols=len(score_extended_column_names))

    # set table style
    score_table.style = document.styles['Light Shading Accent 1']

    hdr_cells = score_table.rows[0].cells

    for i, val in enumerate(score_extended_column_names):
        hdr_cells[i].text = str(val)

    # add a data row for each item
    for i in range(len(score_row_names)):
        cells = score_table.add_row().cells
        # for j in range(len(r2_extended_column_names)):
        for j, val in enumerate(score_extended_column_names):
            if j == 0:
                cells[j].text = score_row_names[i]
            else:
                cells[j].text = str(score_result_df.iat[i, j-1])

                # set the style - text size & bold 
                paragraphs = cells[j].paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(table_font_size) # set the text size

    for row in score_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(table_font_size)

def calculate_mev_values(mape_result_df,training_duration_df, testing_duration_df, config):
    ############################
    # Calculate the MEV
    ############################

    CORRECTNESS_FACTOR = config["general"]["mev_weights"]["correctness"]
    COMPLEXITY_FACTOR = config["general"]["mev_weights"]["complexity"]
    EXPERTISE_FACTOR = config["general"]["mev_weights"]["expertise"]
    RESPONSIVENESS_FACTOR = config["general"]["mev_weights"]["responsiveness"]

    SUM_WEIGHTING_VALUES = CORRECTNESS_FACTOR + COMPLEXITY_FACTOR + EXPERTISE_FACTOR + RESPONSIVENESS_FACTOR

    # 1. merge the average MAPE, training runtime & preiction time results int different subset dataframes
    basic_subset_df = pd.concat([mape_result_df['basic-subset'], training_duration_df['basic-subset'], testing_duration_df['basic-subset']], axis=1)
    extension_df = pd.concat([mape_result_df['extension'], training_duration_df['extension'], testing_duration_df['extension']], axis=1)
    location_df = pd.concat([mape_result_df['location'], training_duration_df['location'], testing_duration_df['location']], axis=1)
    extension_location_df = pd.concat([mape_result_df['extension-location'], training_duration_df['extension-location'], testing_duration_df['extension-location']], axis=1)

    # 2. rename culumn names
    basic_subset_df.columns = ['MAPE', 'Training-time', 'Test-time']
    extension_df.columns = ['MAPE', 'Training-time', 'Test-time']
    location_df.columns = ['MAPE', 'Training-time', 'Test-time']
    extension_location_df.columns = ['MAPE', 'Training-time', 'Test-time']
    # print(extension_location_df)

    # 3. add knowledge levels
    basic_subset_df['K-level'] = [5,5,2,2,2]
    extension_df['K-level'] = [5,5,2,2,2]
    location_df['K-level'] = [5,5,2,2,2]
    extension_location_df['K-level'] = [5,5,2,2,2]
    # print(extension_location_df)

    # 4. normalize the MAPE values
    basic_subset_df['MAPE-norm'] = basic_subset_df['MAPE'] / basic_subset_df['MAPE'].max()
    extension_df['MAPE-norm'] = extension_df['MAPE'] / extension_df['MAPE'].max()
    location_df['MAPE-norm'] = location_df['MAPE'] / location_df['MAPE'].max()
    extension_location_df['MAPE-norm'] = extension_location_df['MAPE'] / extension_location_df['MAPE'].max()
    # print(extension_location_df)

    # 5. normalize the training duration values
    basic_subset_df['Training-norm'] = basic_subset_df['Training-time'] / basic_subset_df['Training-time'].max()
    extension_df['Training-norm'] = extension_df['Training-time'] / extension_df['Training-time'].max()
    location_df['Training-norm'] = location_df['Training-time'] / location_df['Training-time'].max()
    extension_location_df['Training-norm'] = extension_location_df['Training-time'] / extension_location_df['Training-time'].max()
    # print(extension_location_df)

    # 6. normalize the test duration values
    basic_subset_df['Test-norm'] = basic_subset_df['Test-time'] / basic_subset_df['Test-time'].max()
    extension_df['Test-norm'] = extension_df['Test-time'] / extension_df['Test-time'].max()
    location_df['Test-norm'] = location_df['Test-time'] / location_df['Test-time'].max()
    extension_location_df['Test-norm'] = extension_location_df['Test-time'] / extension_location_df['Test-time'].max()
    # print(extension_location_df)

    # 7. normalize knowledge values
    basic_subset_df['K-level-norm'] = basic_subset_df['K-level'] / basic_subset_df['K-level'].max()
    extension_df['K-level-norm'] = extension_df['K-level'] / extension_df['K-level'].max()
    location_df['K-level-norm'] = location_df['K-level'] / location_df['K-level'].max()
    extension_location_df['K-level-norm'] = extension_location_df['K-level'] / extension_location_df['K-level'].max()
    # print(extension_location_df)

    # 8. Weight normed MAPE results
    basic_subset_df['MAPE-weighted'] = basic_subset_df['MAPE-norm'] * CORRECTNESS_FACTOR
    extension_df['MAPE-weighted'] = extension_df['MAPE-norm'] * CORRECTNESS_FACTOR
    location_df['MAPE-weighted'] = location_df['MAPE-norm'] * CORRECTNESS_FACTOR
    extension_location_df['MAPE-weighted'] = extension_location_df['MAPE-norm'] * CORRECTNESS_FACTOR
    # print(extension_location_df)

    # 9. Weight normed training time results
    basic_subset_df['Training-weighted'] = basic_subset_df['Training-norm'] * COMPLEXITY_FACTOR
    extension_df['Training-weighted'] = extension_df['Training-norm'] * COMPLEXITY_FACTOR
    location_df['Training-weighted'] = location_df['Training-norm'] * COMPLEXITY_FACTOR
    extension_location_df['Training-weighted'] = extension_location_df['Training-norm'] * COMPLEXITY_FACTOR
    # print(extension_location_df)

    # 10. Weight normed test time results
    basic_subset_df['Test-weighted'] = basic_subset_df['Test-norm'] * RESPONSIVENESS_FACTOR
    extension_df['Test-weighted'] = extension_df['Test-norm'] * RESPONSIVENESS_FACTOR
    location_df['Test-weighted'] = location_df['Test-norm'] * RESPONSIVENESS_FACTOR
    extension_location_df['Test-weighted'] = extension_location_df['Test-norm'] * RESPONSIVENESS_FACTOR
    # print(extension_location_df)

    # 11. Weight normed knowledge levels
    basic_subset_df['K-level-weighted'] = basic_subset_df['K-level-norm'] * EXPERTISE_FACTOR
    extension_df['K-level-weighted'] = extension_df['K-level-norm'] * EXPERTISE_FACTOR
    location_df['K-level-weighted'] = location_df['K-level-norm'] * EXPERTISE_FACTOR
    extension_location_df['K-level-weighted'] = extension_location_df['K-level-norm'] * EXPERTISE_FACTOR
    # print(extension_location_df)

    # 12. MEV calculation
    basic_subset_df['MEV'] = ((basic_subset_df['MAPE-weighted'] + basic_subset_df['Training-weighted'] + basic_subset_df['Test-weighted'] + basic_subset_df['K-level-weighted']) / SUM_WEIGHTING_VALUES).round(decimals=6)
    extension_df['MEV'] = ((extension_df['MAPE-weighted'] + extension_df['Training-weighted'] + extension_df['Test-weighted'] + extension_df['K-level-weighted']) / SUM_WEIGHTING_VALUES).round(decimals=6)
    location_df['MEV'] = ((location_df['MAPE-weighted'] + location_df['Training-weighted'] + location_df['Test-weighted'] + location_df['K-level-weighted']) / SUM_WEIGHTING_VALUES).round(decimals=6)
    extension_location_df['MEV'] = ((extension_location_df['MAPE-weighted'] + extension_location_df['Training-weighted'] + extension_location_df['Test-weighted'] + extension_location_df['K-level-weighted']) / SUM_WEIGHTING_VALUES).round(decimals=6)
    # print(extension_location_df)

    # 13. merge resulting MEV values into one MEV dataframe
    mev_df = pd.concat([basic_subset_df['MEV'], extension_df['MEV'], location_df['MEV'], extension_location_df['MEV'] ], axis=1)
    mev_df.columns = ['basic-subset', 'extension', 'location', 'extension-location']
    # print(mev_df)

    return mev_df