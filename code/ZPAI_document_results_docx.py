# from asyncio.windows_events import NULL
from logging import NullHandler
from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Pt

from docx.shared import RGBColor


import yaml
from pathlib import Path
#import datetime
from datetime import date

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt

import seaborn as sns

from ZPAI_document_results_utility import load_csv_data, get_min_values, get_max_values, initialize_document, initialize_section, add_score_table, get_results_values, get_duration_results_values, calculate_mev_values
from ZPAI_document_results_utility import add_duration_line_plot, add_results_heading, add_bar_plot, add_stats_table, add_summary_table, add_line_plot, add_box_plot, add_box_plot2, add_simple_score_table
from ZPAI_scratch import plot_feature_performance, plot_training_duration, plot_testing_duration, plot_mev
 
from ZPAI_common_functions import read_yaml



def document_results_docx(const_machine_models: list,
                          NUM_OF_MEASUREMENTS: int,
                          GLOBAL_YAML_SUMMERY_FILE: str, 
                          EXPLICIT_SUMMERY_FILE_PATH: str,
                          config: dict) -> None:

    HEADING_2_LEVELS = 3
    PICTURE_SIZE = 5
    TABLE_FONT_SIZE = 9
    M_DATE = config["general"]["start_date"]

    REPO_PATH = Path(__file__).parents[1]

    NUM_OF_MEASUREMENTS = NUM_OF_MEASUREMENTS

    with open(GLOBAL_YAML_SUMMERY_FILE, 'r') as file: # open the file in append mode
        summery_result_values = yaml.safe_load(file)


    document = initialize_document(summery_result_values=summery_result_values, m_date=M_DATE)


    # iterate through all construction machine models
    for const_machine_model in const_machine_models:

        document = initialize_section(document, summery_result_values, const_machine_model)

        # get values from summery yaml file
        MEASUREMENT_DATE = summery_result_values['measurement_date']
        INPUT_FILE_CREATION_DATE = summery_result_values[const_machine_model]['input_file_creation_date']
        RANDOM_SEED = summery_result_values['random_seed']
        INPUT_FILE_SIZE =  summery_result_values[const_machine_model]['input_file_size']
        INPUT_FILE_NAME =  summery_result_values[const_machine_model]['input_file_name']

        if 'autosklearn_runtime' in summery_result_values:
            AUTOSKLEARN_RUNTIME = summery_result_values['autosklearn_runtime']
            AUTOSKLEARN_LIMIT = summery_result_values['autosklearn_limit']

        # set path variables
        # File path for storing data
        FILE_PATH_DATA = Path('./measurements', const_machine_model, 'data') 

        # File path for storing pictures
        FILE_PATH_PICS = Path('./measurements', const_machine_model, 'pictures')

        # create measurement date directory
        SUB_DIR_DATA = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)

        # construct the file names of the results
        nn_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'nn-results','csv')
        classic_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'classic-results','csv')
        autosklearn_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autosklearn-results', 'csv')
        autogluon_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autogluon-results', 'csv')
        flaml_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'flaml-results', 'csv')
        optuna_result_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'optuna-results', 'csv')

        # init data frames
        nn_results = pd.DataFrame()
        classic_results = pd.DataFrame()
        autosklearn_results = pd.DataFrame()
        autogluon_results = pd.DataFrame()
        flaml_results = pd.DataFrame()
        optuna_results = pd.DataFrame()

        # create result data frame
        r2_result_df = pd.DataFrame()
        rmse_result_df = pd.DataFrame()
        nrmse_result_df = pd.DataFrame()
        iqrrmse_result_df = pd.DataFrame()
        cvrmse_result_df = pd.DataFrame()
        mae_result_df = pd.DataFrame()
        mape_result_df = pd.DataFrame()
        training_duration_df = pd.DataFrame()
        testing_duration_df = pd.DataFrame()
        autosklearn_leaderboard_df = pd.DataFrame()

        # create result data frame for scatter plot
        scatter_r2_result_df = pd.DataFrame()
        scatter_r2_rmse_result_df = pd.DataFrame()
        scatter_r2_nrmse_result_df = pd.DataFrame()
        scatter_r2_iqrrmse_result_df = pd.DataFrame()
        scatter_r2_cvrmse_result_df = pd.DataFrame()
        scatter_r2_mae_result_df = pd.DataFrame()
        scatter_r2_mape_result_df = pd.DataFrame()
        scatter_training_duration_result_df = pd.DataFrame()
        scatter_testing_duration_result_df = pd.DataFrame()

        # get the results for each measurement and add it to the previous one
        for measurement in range(NUM_OF_MEASUREMENTS):

            # get nn results
            NN_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), nn_result_filename)
            if NN_FILE_PATH.is_file():
                nn_results = load_csv_data(NN_FILE_PATH)

            # get classical results
            CLASSIC_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), classic_result_filename)
            if CLASSIC_FILE_PATH.is_file():
                classic_results = load_csv_data(CLASSIC_FILE_PATH)

            # get autosklearn results
            AUTOSKLEARN_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), autosklearn_result_filename)
            if AUTOSKLEARN_FILE_PATH.is_file():
                autosklearn_results = load_csv_data(AUTOSKLEARN_FILE_PATH)


            # get autogluon results
            AUTOGLUON_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), autogluon_result_filename)
            if AUTOGLUON_FILE_PATH.is_file():
                autogluon_results = load_csv_data(AUTOGLUON_FILE_PATH)

            # get flaml results
            FLAML_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), flaml_result_filename)
            if FLAML_FILE_PATH.is_file():
                flaml_results = load_csv_data(FLAML_FILE_PATH)

            # get optuna results
            OPTUNA_FILE_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'data', SUB_DIR_DATA, str(measurement + 1), optuna_result_filename)
            if OPTUNA_FILE_PATH.is_file():
                optuna_results = load_csv_data(OPTUNA_FILE_PATH)

            # get the result values
            temp_r2_result_df = get_results_values("R2", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_mae_result_df = get_results_values("MAE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_mape_result_df = get_results_values("MAPE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_rmse_result_df = get_results_values("RMSE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_nrmse_result_df = get_results_values("N-RMSE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_iqrrmse_result_df = get_results_values("IQR-RMSE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_cvrmse_result_df = get_results_values("CV-RMSE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_training_duration_result_df = get_duration_results_values("Training-Duration", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
            temp_testing_duration_result_df = get_duration_results_values("Test-Duration", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)


            # init the dataframe for the first time
            if measurement == 0:
                r2_result_df = temp_r2_result_df
                mae_result_df = temp_mae_result_df
                mape_result_df = temp_mape_result_df
                rmse_result_df = temp_rmse_result_df
                nrmse_result_df = temp_nrmse_result_df
                iqrrmse_result_df = temp_iqrrmse_result_df
                cvrmse_result_df = temp_cvrmse_result_df
                training_duration_df = temp_training_duration_result_df
                testing_duration_df = temp_testing_duration_result_df

                scatter_r2_result_df = temp_r2_result_df
                scatter_r2_rmse_result_df = temp_rmse_result_df
                scatter_r2_nrmse_result_df = temp_nrmse_result_df
                scatter_r2_iqrrmse_result_df = temp_iqrrmse_result_df
                scatter_r2_cvrmse_result_df = temp_cvrmse_result_df
                scatter_r2_mae_result_df = temp_mae_result_df
                scatter_r2_mape_result_df = temp_mape_result_df
                scatter_training_duration_result_df = temp_training_duration_result_df
                scatter_testing_duration_result_df = temp_testing_duration_result_df


            # add values to existing dataframe
            else: 
                r2_result_df = r2_result_df + temp_r2_result_df
                mae_result_df = mae_result_df + temp_mae_result_df
                mape_result_df = mape_result_df + temp_mape_result_df
                rmse_result_df = rmse_result_df + temp_rmse_result_df
                nrmse_result_df = nrmse_result_df + temp_nrmse_result_df
                iqrrmse_result_df = iqrrmse_result_df + temp_iqrrmse_result_df
                cvrmse_result_df = cvrmse_result_df + temp_cvrmse_result_df
                training_duration_df = training_duration_df + temp_training_duration_result_df
                testing_duration_df = testing_duration_df + temp_testing_duration_result_df

                scatter_r2_result_df = pd.concat([scatter_r2_result_df, temp_r2_result_df])
                scatter_r2_rmse_result_df = pd.concat([scatter_r2_rmse_result_df, temp_rmse_result_df])
                scatter_r2_nrmse_result_df = pd.concat([scatter_r2_nrmse_result_df, temp_nrmse_result_df])
                scatter_r2_iqrrmse_result_df = pd.concat([scatter_r2_iqrrmse_result_df, temp_iqrrmse_result_df])
                scatter_r2_cvrmse_result_df = pd.concat([scatter_r2_cvrmse_result_df, temp_cvrmse_result_df])
                scatter_r2_mae_result_df = pd.concat([scatter_r2_mae_result_df, temp_mae_result_df])
                scatter_r2_mape_result_df = pd.concat([scatter_r2_mape_result_df, temp_mape_result_df])
                scatter_training_duration_result_df = pd.concat([scatter_training_duration_result_df, temp_training_duration_result_df])
                scatter_testing_duration_result_df = pd.concat([scatter_testing_duration_result_df, temp_testing_duration_result_df])
            
            # print('INSIDE: \n', scatter_training_duration_result_df)

        # print('\nOUTSIDE before average: \n', scatter_r2_mape_result_df)

        # calculate the averages for each value
        r2_result_df = (r2_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        mae_result_df = (mae_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        mape_result_df = (mape_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        rmse_result_df = (rmse_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        nrmse_result_df = (nrmse_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        iqrrmse_result_df = (iqrrmse_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        cvrmse_result_df = (cvrmse_result_df / NUM_OF_MEASUREMENTS).round(decimals=4)
        training_duration_df = (training_duration_df / NUM_OF_MEASUREMENTS).round(decimals=1)
        testing_duration_df = (testing_duration_df / NUM_OF_MEASUREMENTS).round(decimals=10)

        # print('\nOUTSIDE after average: \n', training_duration_df)

        # print(scatter_r2_mape_result_df)

        ################################
        # Display R2 results
        ################################

        add_results_heading(document, f"R2 results for {const_machine_model}", HEADING_2_LEVELS)
        # r2_result_df = get_results_values("R2", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
        add_score_table(document, r2_result_df, "R2", TABLE_FONT_SIZE)
        add_bar_plot(document, r2_result_df, "R2", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        # add_line_plot(document, r2_result_df, "R2", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        document.add_page_break()

        #######################
        # Display MAE results
        #######################

        add_results_heading(document, f"MAE results for {const_machine_model}", HEADING_2_LEVELS)
        # mae_result_df = get_results_values("MAE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
        add_score_table(document, mae_result_df, "MAE", TABLE_FONT_SIZE)
        add_bar_plot(document, mae_result_df, "MAE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        # add_line_plot(document, mae_result_df, "MAE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        document.add_page_break()

        #######################
        # Display MAPE results
        #######################
        add_results_heading(document, f"MAPE results for {const_machine_model}", HEADING_2_LEVELS)
        # mape_result_df = get_results_values("MAPE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
        add_score_table(document, mape_result_df, "MAPE", TABLE_FONT_SIZE)
        add_bar_plot(document, mape_result_df, "MAPE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        # add_line_plot(document, mape_result_df, "MAPE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)

        #######################
        # Display scatter plot for MAPE results
        #######################
        # construct file / path to save the scatter plot
        sub_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'MAPE-scatter-plot', 'png')
        pdf_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'MAPE-scatter-plot', 'pdf')
        CHART_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
        CHART_PDF_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, pdf_filename)

        # calculate min and max values for scaling the plot
        score_col_name_max, score_row_name_max, max_value =  get_max_values(scatter_r2_mape_result_df) # get max value for scaling the plot
        score_col_name_min, score_row_name_min, min_value =  get_min_values(scatter_r2_mape_result_df) # get min value for scaling the plot

        # 1. add column name to first column
        sc1 = scatter_r2_mape_result_df.rename_axis('names')
        # 2. groupby the name of the algorithms and gather the results in a list
        sc_mape = sc1.groupby('names').agg(lambda x: list(x))

        plot_feature_performance(sc_mape, "MAPE", CHART_PATH, CHART_PDF_PATH, max_value, min_value, const_machine_model, document, PICTURE_SIZE)

        document.add_page_break()

        #######################
        # Display RMSE results
        #######################

        add_results_heading(document, f"RMSE results for {const_machine_model}", HEADING_2_LEVELS)
        # rmse_result_df = get_results_values("RMSE", classic_results, nn_results, autosklearn_results, autogluon_results, flaml_results, optuna_results)
        add_score_table(document, rmse_result_df, "RMSE", TABLE_FONT_SIZE)
        add_bar_plot(document, rmse_result_df, "RMSE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)
        # add_line_plot(document, rmse_result_df, "RMSE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)

        document.add_page_break()

        ################################
        # Display training duration
        ################################

        add_results_heading(document, f"Training run-time of each method and feature subset in seconds for {const_machine_model}", HEADING_2_LEVELS)

        # ----------------
        # construct table
        # ----------------
        duration_column_names = list(training_duration_df.columns.values)
        duration_row_names = list(training_duration_df.index.values)
        duration_column_names.insert(0, ' ')
        duration_table = document.add_table(rows=1, cols=len(duration_column_names))

        # set table style
        duration_table.style = document.styles['Light Shading Accent 1']

        hdr_cells = duration_table.rows[0].cells

        for i, val in enumerate(duration_column_names):
            hdr_cells[i].text = str(val)

        # add a data row for each item
        for i in range(len(duration_row_names)):
            cells = duration_table.add_row().cells
            for j in range(len(duration_column_names)):
                if j == 0:
                    cells[j].text = duration_row_names[i]
                else:
                    cells[j].text = str(training_duration_df.iat[i, j-1])

        for row in duration_table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(TABLE_FONT_SIZE)

        # Add duration line plot
        # add_duration_line_plot(document, training_duration_df, "Training-Duration", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)

        #######################
        # Display scatter plot for training duration
        #######################
        # construct file / path to save the scatter plot
        sub_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'Training-duration-scatter-plot', 'png')
        pdf_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'Training-duration-scatter-plot', 'pdf')
        CHART_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
        CHART_PDF_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, pdf_filename)

        # calculate min and max values for scaling the plot
        score_col_name_max, score_row_name_max, max_value =  get_max_values(scatter_training_duration_result_df) # get max value for scaling the plot
        score_col_name_min, score_row_name_min, min_value =  get_min_values(scatter_training_duration_result_df) # get min value for scaling the plot

        # 1. add column name to first column
        sc1 = scatter_training_duration_result_df.rename_axis('names')
        # 2. groupby the name of the algorithms and gather the results in a list
        sc_training = sc1.groupby('names').agg(lambda x: list(x))
        # print(sc)
        plot_training_duration(sc_training, "MAPE", CHART_PATH, CHART_PDF_PATH, max_value, min_value, const_machine_model, document, PICTURE_SIZE)

        # document.add_page_break()

        ################################
        # Display test duration
        ################################

        add_results_heading(document, f"Prediction time of each method and feature subset in seconds for {const_machine_model}", HEADING_2_LEVELS)

        # ----------------
        # construct table
        # ----------------
        duration_column_names = list(testing_duration_df.columns.values)
        duration_row_names = list(testing_duration_df.index.values)
        duration_column_names.insert(0, ' ')
        duration_table = document.add_table(rows=1, cols=len(duration_column_names))

        # set table style
        duration_table.style = document.styles['Light Shading Accent 1']

        hdr_cells = duration_table.rows[0].cells

        for i, val in enumerate(duration_column_names):
            hdr_cells[i].text = str(val)

        # add a data row for each item
        for i in range(len(duration_row_names)):
            cells = duration_table.add_row().cells
            for j in range(len(duration_column_names)):
                if j == 0:
                    cells[j].text = duration_row_names[i]
                else:
                    cells[j].text = str(testing_duration_df.iat[i, j-1])

        for row in duration_table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(TABLE_FONT_SIZE)

        # Add duration line plot
        # add_duration_line_plot(document, testing_duration_df, "Prediction-Duration", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)

        #######################
        # Display scatter plot for testing duration
        #######################
        # construct file / path to save the scatter plot
        sub_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'Prediction-duration-scatter-plot', 'png')
        pdf_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'Prediction-duration-scatter-plot', 'pdf')
        CHART_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
        CHART_PDF_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, pdf_filename)


        # calculate min and max values for scaling the plot
        score_col_name_max, score_row_name_max, max_value =  get_max_values(scatter_testing_duration_result_df) # get max value for scaling the plot
        score_col_name_min, score_row_name_min, min_value =  get_min_values(scatter_testing_duration_result_df) # get min value for scaling the plot

        # 1. add column name to first column
        sc1 = scatter_testing_duration_result_df.rename_axis('names')
        # 2. groupby the name of the algorithms and gather the results in a list
        sc_test = sc1.groupby('names').agg(lambda x: list(x))
        # print(sc)
        # print(sc.iloc[0][0])
        plot_testing_duration(sc_test, "MAPE", CHART_PATH, CHART_PDF_PATH, max_value, min_value, const_machine_model, document, PICTURE_SIZE)

        document.add_page_break()

        # ############################
        # # Calculate the MEV
        # ############################

        mev_df = calculate_mev_values(mape_result_df,training_duration_df, testing_duration_df, config)

         # construct file / path to save the scatter plot
        sub_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'MEV-scatter-plot', 'png')
        pdf_filename = "{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,f'MEV-scatter-plot', 'pdf')
        CHART_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, filename)
        CHART_PDF_PATH = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures',sub_dir, pdf_filename)


        # calculate min and max values for scaling the plot
        score_col_name_max, score_row_name_max, max_value =  get_max_values(mev_df) # get max value for scaling the plot
        score_col_name_min, score_row_name_min, min_value =  get_min_values(mev_df) # get min value for scaling the plot

        add_results_heading(document, f"MEV results for all feature sets.", HEADING_2_LEVELS)
        add_score_table(document, mev_df, 'MEV', TABLE_FONT_SIZE)
        best_mev_feature_set, best_mev_method, best_mev_value = get_min_values(mev_df)

        # print(mev_df)
        mev_df = mev_df.reindex(['autogluon', 'autosklearn', 'classic', 'flaml', 'nn'])
        # print(mev_df)

        plot_mev(mev_df, CHART_PATH, CHART_PDF_PATH, max_value, min_value, const_machine_model, document, PICTURE_SIZE)

        #######################
        # calculate standart deviation for the MAPA measurements 
        # for the best feature-set combination
        #######################

        # extract feature-set combination column with best MAPE score
        best_res_mape = sc_mape.loc[:, [best_mev_feature_set]]

        # create new dataframe
        std_dataframe = best_res_mape.copy()

        # calculate the standard deviation and stor it in std_dataframe
        for index, row in best_res_mape.iterrows():
            std_dataframe.loc[index, best_mev_feature_set] = np.std(row[best_mev_feature_set]).round(decimals=8)

        # print(std_dataframe)
        add_results_heading(document, f"Standard deviation for MAPE results for {best_mev_feature_set} feature set.", HEADING_2_LEVELS)
        add_simple_score_table(document, std_dataframe, TABLE_FONT_SIZE)
       
        ######################
        # Display MAPE box plot for best feature-set combination
        #####################
        add_results_heading(document, f"Boxplot for MAPE results for {best_mev_feature_set} feature set.", HEADING_2_LEVELS)
        add_box_plot2(document, best_res_mape, best_mev_feature_set, "MAPE", MEASUREMENT_DATE, INPUT_FILE_CREATION_DATE, const_machine_model, REPO_PATH, PICTURE_SIZE)

        #######################
        # Calculate standart deviation for the training runtime measurements 
        # for the best feature-set combination
        #######################

        # extract feature-set combination column with best MAPE score
        best_res_training = sc_training.loc[:, [best_mev_feature_set]]

        # create new dataframe
        std_dataframe = best_res_training.copy()

        # calculate the standard deviation and stor it in std_dataframe
        for index, row in best_res_training.iterrows():
            std_dataframe.loc[index, best_mev_feature_set] = np.std(row[best_mev_feature_set]).round(decimals=8)

        # print(std_dataframe)
        add_results_heading(document, f"Standard deviation for training time for {best_mev_feature_set} feature set.", HEADING_2_LEVELS)
        add_simple_score_table(document, std_dataframe, TABLE_FONT_SIZE)

        #######################
        # Calculate standart deviation for prediction 
        # for the best feature-set combination
        #######################

        # extract feature-set combination column with best MAPE score
        best_res_test = sc_test.loc[:, [best_mev_feature_set]]

        # create new dataframe
        std_dataframe = best_res_test.copy()

        # calculate the standard deviation and stor it in std_dataframe
        for index, row in best_res_test.iterrows():
            std_dataframe.loc[index, best_mev_feature_set] = np.std(row[best_mev_feature_set]).round(decimals=12)

        # print(std_dataframe)
        add_results_heading(document, f"Standard deviation for prediction time for {best_mev_feature_set} feature set.", HEADING_2_LEVELS)
        add_simple_score_table(document, std_dataframe, TABLE_FONT_SIZE)

        ################################################
        # Display best classical model for best R2-score
        ################################################

        if not classic_results.empty:

            add_results_heading(document, f"Best classical model for {const_machine_model}", HEADING_2_LEVELS)

            # extract the column name with the maximal R2 score
            max_col_values = classic_results.loc[ 'Test-R2', :]
            max_col_values = max_col_values.to_frame()
            max_col_values['Test-R2'] = pd.to_numeric(max_col_values['Test-R2'])
            max_column = max_col_values.idxmax().iloc[0]

            # extract/copy the column with the best R2-score
            best_classic_model = classic_results[max_column].copy()

            # print(best_classic_model.loc['final-model'])

            # construct table with the following column names
            column_names = ['Best classical model']

            table = document.add_table(rows=1, cols=len(column_names))

            # set table style
            table.style = document.styles['Light Shading Accent 1']

            hdr_cells = table.rows[0].cells

            # build the header row
            for i, val in enumerate(column_names):
                hdr_cells[i].text = str(val)

            # add a cell to the table & fill it with the data
            cells = table.add_row().cells
            for j in range(len(column_names)):
                cells[j].text = str(best_classic_model.loc['final-model'])

            # set table font size
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(TABLE_FONT_SIZE)

        #############################################
        # Display best NN params for best R2-score
        #############################################

        if not nn_results.empty:

            add_results_heading(document, f"Best NN parameter set for {const_machine_model}", HEADING_2_LEVELS)

            # extract the column name with the maximal R2 score
            max_col_values = nn_results.loc[ 'Test-R2', :]
            max_col_values = max_col_values.to_frame()
            max_col_values['Test-R2'] = pd.to_numeric(max_col_values['Test-R2'])
            max_column = max_col_values.idxmax().iloc[0]

            # extract/copy the column with the best R2-score
            best_nn_val = nn_results[max_column].copy()

            # construct table with the following column names
            column_names = ['Activation', 'Hidden-layer-size', 'Learning rate',  'Solver']

            table = document.add_table(rows=1, cols=len(column_names))

            # set table style
            table.style = document.styles['Light Shading Accent 1']

            hdr_cells = table.rows[0].cells

            # build the header row
            for i, val in enumerate(column_names):
                hdr_cells[i].text = str(val)

            # add a cell to the table & fill it with the data
            cells = table.add_row().cells
            for j in range(len(column_names)):
                cells[j].text = str(best_nn_val.iloc[j+9])

            # set table font size
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(TABLE_FONT_SIZE)



        #################################
        # Display autosklearn leaderboard - ToDo: update path
        #################################
        rmse_col_name_min, rmse_row_name_min, min_value =  get_min_values(rmse_result_df)

        # File path for stored leaderboards
        leaderboards_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        best_leaderboard_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autosklearn-leaderboard',rmse_col_name_min, 'csv')
        FILE_PATH_LEADERBOARD = Path(REPO_PATH, 'measurements', const_machine_model, 'data', leaderboards_dir, best_leaderboard_file)
        FILE_PATH_LEADERBOARD_STR = str(Path(REPO_PATH, 'measurements', const_machine_model, 'data', leaderboards_dir, best_leaderboard_file))

        print(FILE_PATH_LEADERBOARD_STR)

        # check if leaderboard file for autosklearn exists
        if FILE_PATH_LEADERBOARD.is_file():

            
            add_results_heading(document, f"Autosklearn leaderboard for features: {rmse_col_name_min}", HEADING_2_LEVELS)

            # load the appropriate leaderboard
            leaderboard_results = load_csv_data(FILE_PATH_LEADERBOARD_STR)

            # construct table with the following column names
            extracted_column_names = ['Rank', 'Method', 'Ensemble weight']

            # get the number of rows
            row_number = list(leaderboard_results.index)

            table = document.add_table(rows=1, cols=len(extracted_column_names))

            # set table style
            table.style = document.styles['Light Shading Accent 1']

            hdr_cells = table.rows[0].cells

            # build the header row
            for i, val in enumerate(extracted_column_names):
                hdr_cells[i].text = str(val)

            # add a data row for each item
            for i in range(len(row_number)):
                cells = table.add_row().cells
                for j in range(len(extracted_column_names)):
                    # cells[j].text = str(leaderboard_results.iat[i, j])
                    if j == 0:
                        cells[j].text = str(leaderboard_results.iat[i, j])
                    elif j == 1:
                        cells[j].text = str(leaderboard_results.iat[i, j+1])
                    elif j == 2:
                        cells[j].text = str(leaderboard_results.iat[i, j-1])

            # set table font size
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(TABLE_FONT_SIZE)

        #################################
        # Display autogluon leaderboard - ToDo: update path
        #################################

        # File path for stored leaderboards
        leaderboards_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        best_leaderboard_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autogluon-leaderboard',rmse_col_name_min, 'csv')
        FILE_PATH_LEADERBOARD = Path(REPO_PATH, 'measurements', const_machine_model, 'data', leaderboards_dir, best_leaderboard_file)
        FILE_PATH_LEADERBOARD_STR = str(Path(REPO_PATH, 'measurements', const_machine_model, 'data', leaderboards_dir, best_leaderboard_file))

        # check if leaderboard file for autogluon exists
        if FILE_PATH_LEADERBOARD.is_file():

            add_results_heading(document, f"Autogluon leaderboard for features: {rmse_col_name_min}", HEADING_2_LEVELS)

            # load the appropriate leaderboard
            leaderboard_results = load_csv_data(FILE_PATH_LEADERBOARD_STR)

            # construct table with the following column names
            extracted_column_names = ['Rank', 'Score test', 'Score Value']

            # get the number of rows
            row_number = list(leaderboard_results.index)

            table = document.add_table(rows=1, cols=len(extracted_column_names))

            # set table style
            table.style = document.styles['Light Shading Accent 1']

            hdr_cells = table.rows[0].cells

            # build the header row
            for i, val in enumerate(extracted_column_names):
                hdr_cells[i].text = str(val)

            # add a data row for each item
            for i in range(len(row_number)):
                cells = table.add_row().cells
                for j in range(len(extracted_column_names)):
                    # cells[j].text = str(leaderboard_results.iat[i, j])
                    if j == 0:
                        cells[j].text = str(leaderboard_results.iat[i, j])
                    elif j == 1:
                        cells[j].text = str(leaderboard_results.iat[i, j+1])
                    elif j == 2:
                        cells[j].text = str(leaderboard_results.iat[i, j-1])

            # set table font size
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(TABLE_FONT_SIZE)


        ############################
        # Display test result chars
        ############################
        # add_results_heading(document, f"Test results for {const_machine_model}", HEADING_2_LEVELS)

        # # Calculate the min value for RMSE and the corresponding column and index names
        # rmse_location, rmse_method, rmse_min_value = get_min_values(rmse_result_df)
        # rmse_location, rmse_method, rmse_min_value

        # # File path for storing pictures
        # best_result_picture_dir = "{}-{}-{}-{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE)
        # FILE_PATH_PICS = Path(REPO_PATH, 'measurements', const_machine_model, 'pictures', best_result_picture_dir)
        # if rmse_method == 'autosklearn':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autosklearn-test-result',rmse_location, 'png')
        # elif rmse_method == 'autogluon':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'autogluon-test-result',rmse_location, 'png')
        # elif rmse_method == 'nn':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'nn-test-result',rmse_location, 'png')
        # elif rmse_method == 'classic':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'classic-test-result',rmse_location, 'png')
        # elif rmse_method == 'flaml':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'flaml-test-result',rmse_location, 'png')
        # elif rmse_method == 'optuna':
        #     best_result_picture_file = "{}-{}-{}-{}-{}-{}.{}".format(MEASUREMENT_DATE, const_machine_model,'final',INPUT_FILE_CREATION_DATE,'optuna-test-result',rmse_location, 'png')
        # else:
        #     print('Failure in displaying the best result')

        #  # get the plots from the first '1' measurement
        # FILE_PATH_PICS = str(Path(REPO_PATH, 'measurements', const_machine_model, 'pictures', best_result_picture_dir, '1', best_result_picture_file))

        # listOfImageNames = [FILE_PATH_PICS]
        # for imageName in listOfImageNames:
        #     document.add_picture(str(imageName), width=Inches(PICTURE_SIZE))

        #######################
        # Display best results
        #######################

        add_results_heading(document, f"Best results for {const_machine_model}", HEADING_2_LEVELS)

        # create resulting data frame
        result_df_index = ['Feature-Set', 'Method', 'Value']
        result_df = pd.DataFrame(index=result_df_index)

        result_df['R2-Score'] = get_max_values(r2_result_df)
        result_df['MAPE'] = get_min_values(mape_result_df)
        result_df['MAE'] = get_min_values(mae_result_df)
        result_df['RMSE'] = get_min_values(rmse_result_df)
        result_df['N-RMSE'] = get_min_values(nrmse_result_df)
        result_df['IQR-RMSE'] = get_min_values(iqrrmse_result_df)
        result_df['CV-RMSE'] = get_min_values(cvrmse_result_df)

        # store resulting dataframe
        filename = "{}-{}-{}.{}".format(M_DATE, const_machine_model, 'best-results','csv')
        RESULT_SUMMERY_FILE = Path(EXPLICIT_SUMMERY_FILE_PATH, filename)
        result_df.to_csv(str(RESULT_SUMMERY_FILE))

        # construct table
        extended_column_names = list(result_df.columns.values)
        row_names = list(result_df.index.values)

        extended_column_names.insert(0, ' ')

        # table = document.add_table(rows=len(nn_row_names), cols=len(nn_extended_column_names))
        table = document.add_table(rows=1, cols=len(extended_column_names))

        # set table style
        table.style = document.styles['Light Shading Accent 1']

        hdr_cells = table.rows[0].cells

        for i, val in enumerate(extended_column_names):
            hdr_cells[i].text = str(val)

        # add a data row for each item
        for i in range(len(row_names)):
            cells = table.add_row().cells
            for j in range(len(extended_column_names)):
                if j == 0:
                    cells[j].text = row_names[i]
                else:
                    cells[j].text = str(result_df.iat[i, j-1])

        # set table font size
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(TABLE_FONT_SIZE)


    ##########################
    # Build the summery table
    ##########################

    add_results_heading(document, "Summery of the best results for all data sets", 1)

    # create result data frame
    best_result_df = pd.DataFrame()

    # iterate through all construction machine models
    for const_machine_model in const_machine_models:
        filename = "{}-{}-{}.{}".format(M_DATE, const_machine_model, 'best-results','csv')
        RESULT_SUMMERY_FILE = Path(EXPLICIT_SUMMERY_FILE_PATH, filename)

        results = load_csv_data(RESULT_SUMMERY_FILE)

        # Extract Values row for results.
        value_results = results.loc[ 'Value' , : ]
        best_result_df = pd.concat([best_result_df, pd.DataFrame.from_records([value_results])])
        best_result_df = best_result_df.rename(index={0: str(const_machine_model)})

    
    #------------------------
    # get the min and max values
    #------------------------
    add_summary_table(document, best_result_df, summery_result_values, TABLE_FONT_SIZE)


    ##########################
    # Build the stats table
    ##########################

    add_results_heading(document, "Statistics for all data sets", 1)
    add_stats_table(document, best_result_df, summery_result_values, TABLE_FONT_SIZE)

    # save best_result_df
    filename = "{}-{}.{}".format(M_DATE, 'best-results-for-all-machines','csv')
    RESULT_SUMMERY_FILE = Path(EXPLICIT_SUMMERY_FILE_PATH, filename)
    best_result_df.to_csv(str(RESULT_SUMMERY_FILE))

    ######################
    # Save the docx file
    ######################
    SUMMERY_FILE_PATH = Path(REPO_PATH, 'measurements', 'summery', M_DATE)
    filename = "{}-{}.{}".format(M_DATE,'measurement-documentation','docx')
    DOCX_FILE = Path(SUMMERY_FILE_PATH, filename)
    document.save(str(DOCX_FILE))

###############################
# FOR MANUAL DOCUMENT CREATION
###############################
# const_machine_models = ['Caterpillar-308', 'Caterpillar-320', 'Caterpillar-323', 'Caterpillar-329', 'Caterpillar-330', 'Caterpillar-336', 'Caterpillar-950', 'Caterpillar-966', 'Caterpillar-D6', 'Caterpillar-M318']
const_machine_models = ['merged-files']
# const_machine_models = ['Caterpillar-950', 'Caterpillar-966']
# const_machine_models = ['Caterpillar-323']
# get cuttent date
today = date.today()
# # YY-mm-dd
m_date = today.strftime("%Y-%m-%d")

# DELETE after testing
m_date = '2023-02-16'

# create summery yaml file
# File path within the summery directory for each measurement
EXPLICIT_SUMMERY_FILE_PATH = Path('./measurements', 'summery', m_date)

filename = "{}-{}.{}".format(m_date,'summery','yml')
GLOBAL_YAML_SUMMERY_FILE = Path(EXPLICIT_SUMMERY_FILE_PATH, filename)

######################################
# LOAD CONFIGURATIONS
######################################
REPO_PATH = Path(__file__).parents[1]

# Load configuration files
general_conf = read_yaml(REPO_PATH / 'conf/general_config.yml')
model_conf = read_yaml(REPO_PATH / 'conf/model_config.yml')
autosklearn_conf = read_yaml(REPO_PATH / 'conf/auto_sklearn_config.yml')

# Create global configuration file
CFG = dict()
CFG["general"] = general_conf
CFG["model"] = model_conf
CFG["autosklearn"] = autosklearn_conf

# CFG["general"]["start_date"] = get_current_date()
CFG["general"]["start_date"] = m_date

NUM_OF_MEASUREMENTS = 5

# document_results_docx(const_machine_models,m_date, GLOBAL_YAML_SUMMERY_FILE, EXPLICIT_SUMMERY_FILE_PATH)
document_results_docx(const_machine_models,
                              NUM_OF_MEASUREMENTS = NUM_OF_MEASUREMENTS,
                              GLOBAL_YAML_SUMMERY_FILE = GLOBAL_YAML_SUMMERY_FILE, 
                              EXPLICIT_SUMMERY_FILE_PATH = EXPLICIT_SUMMERY_FILE_PATH, 
                              config = CFG)
